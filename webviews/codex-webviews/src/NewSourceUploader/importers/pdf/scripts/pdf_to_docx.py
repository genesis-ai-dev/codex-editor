#!/usr/bin/env python3
"""
PDF to DOCX Converter - Hybrid Approach

Tries multiple conversion methods in order of quality:
1. LibreOffice headless (best layout preservation, free)
2. pdf2docx library (good for most PDFs)
3. Rich text extraction fallback (preserves content when others fail)

Preserves fonts, sizes, colors, images, line breaks, page breaks,
tables, headers/footers, and document structure.
"""

import sys
import os
import json
import io
import tempfile
import shutil
import subprocess
from pathlib import Path

# Check for required libraries
try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

try:
    from pdf2docx import Converter
    HAS_PDF2DOCX = True
except ImportError:
    HAS_PDF2DOCX = False

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Emu, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False

try:
    from PIL import Image
    HAS_PIL = True
except ImportError:
    HAS_PIL = False


def find_libreoffice() -> str | None:
    """Find LibreOffice executable on the system."""
    if sys.platform == 'win32':
        # Common Windows paths
        possible_paths = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            os.path.expandvars(r"%PROGRAMFILES%\LibreOffice\program\soffice.exe"),
            os.path.expandvars(r"%PROGRAMFILES(X86)%\LibreOffice\program\soffice.exe"),
        ]
        for path in possible_paths:
            if os.path.exists(path):
                return path
        # Try PATH
        try:
            result = subprocess.run(['where', 'soffice'], capture_output=True, text=True)
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout.strip().split('\n')[0]
        except:
            pass
    else:
        # macOS / Linux
        possible_paths = [
            '/usr/bin/soffice',
            '/usr/bin/libreoffice',
            '/Applications/LibreOffice.app/Contents/MacOS/soffice',
            '/opt/libreoffice/program/soffice',
        ]
        for path in possible_paths:
            if os.path.exists(path):
                return path
        # Try PATH
        try:
            result = subprocess.run(['which', 'soffice'], capture_output=True, text=True)
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout.strip()
        except:
            pass
        try:
            result = subprocess.run(['which', 'libreoffice'], capture_output=True, text=True)
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout.strip()
        except:
            pass
    
    return None


def convert_with_libreoffice(pdf_path: str, output_path: str) -> dict:
    """
    Convert PDF to DOCX using LibreOffice headless mode.
    This method provides the best layout preservation for most PDFs.
    """
    soffice = find_libreoffice()
    if not soffice:
        return {
            "success": False,
            "error": "LibreOffice not found. Install from https://www.libreoffice.org/",
            "method": "libreoffice"
        }
    
    try:
        print(json.dumps({"info": "Converting with LibreOffice (best quality)..."}), file=sys.stderr)
        
        # Create temp directory for output
        temp_dir = tempfile.mkdtemp(prefix="lo_convert_")
        
        try:
            # LibreOffice command for PDF to DOCX conversion
            # --infilter specifies PDF import filter
            # --convert-to specifies output format
            cmd = [
                soffice,
                "--headless",
                "--infilter=writer_pdf_import",
                "--convert-to", "docx:Office Open XML Text",
                "--outdir", temp_dir,
                pdf_path
            ]
            
            print(json.dumps({"info": f"Running: {' '.join(cmd[:4])}..."}), file=sys.stderr)
            
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=600  # 10 minute timeout for large PDFs
            )
            
            if result.returncode != 0:
                error_msg = result.stderr or result.stdout or "Unknown error"
                return {
                    "success": False,
                    "error": f"LibreOffice conversion failed: {error_msg}",
                    "method": "libreoffice"
                }
            
            # Find the output file
            pdf_basename = os.path.splitext(os.path.basename(pdf_path))[0]
            temp_output = os.path.join(temp_dir, f"{pdf_basename}.docx")
            
            if not os.path.exists(temp_output):
                # Try alternative naming
                for f in os.listdir(temp_dir):
                    if f.endswith('.docx'):
                        temp_output = os.path.join(temp_dir, f)
                        break
            
            if not os.path.exists(temp_output):
                return {
                    "success": False,
                    "error": "LibreOffice did not create output file",
                    "method": "libreoffice"
                }
            
            # Check file size
            file_size = os.path.getsize(temp_output)
            if file_size == 0:
                return {
                    "success": False,
                    "error": "LibreOffice created empty output file",
                    "method": "libreoffice"
                }
            
            # Move to final location
            shutil.move(temp_output, output_path)
            
            print(json.dumps({"info": f"LibreOffice conversion successful ({file_size} bytes)"}), file=sys.stderr)
            
            return {
                "success": True,
                "outputPath": output_path,
                "method": "libreoffice"
            }
            
        finally:
            shutil.rmtree(temp_dir, ignore_errors=True)
            
    except subprocess.TimeoutExpired:
        return {
            "success": False,
            "error": "LibreOffice conversion timed out after 10 minutes",
            "method": "libreoffice"
        }
    except Exception as e:
        return {
            "success": False,
            "error": f"LibreOffice error: {str(e)}",
            "method": "libreoffice"
        }


def convert_with_pdf2docx(pdf_path: str, output_path: str) -> dict:
    """
    Convert PDF to DOCX using pdf2docx library.
    Good for most text-based PDFs.
    """
    if not HAS_PDF2DOCX:
        return {
            "success": False,
            "error": "pdf2docx not installed. Install with: pip install pdf2docx",
            "method": "pdf2docx"
        }
    
    try:
        print(json.dumps({"info": "Converting with pdf2docx..."}), file=sys.stderr)
        
        cv = Converter(pdf_path)
        cv.convert(output_path, start=0, end=None)
        cv.close()
        
        # Verify output
        if not os.path.exists(output_path):
            return {
                "success": False,
                "error": "pdf2docx did not create output file",
                "method": "pdf2docx"
            }
        
        file_size = os.path.getsize(output_path)
        if file_size == 0:
            return {
                "success": False,
                "error": "pdf2docx created empty output file",
                "method": "pdf2docx"
            }
        
        print(json.dumps({"info": f"pdf2docx conversion successful ({file_size} bytes)"}), file=sys.stderr)
        
        return {
            "success": True,
            "outputPath": output_path,
            "method": "pdf2docx"
        }
        
    except Exception as e:
        error_msg = str(e)
        # Check for known issues
        is_recoverable = any(x in error_msg.lower() for x in [
            "pixmap must be grayscale or rgb",
            "code=4",
            "colorspace",
            "cmyk"
        ])
        
        return {
            "success": False,
            "error": error_msg,
            "method": "pdf2docx",
            "recoverable": is_recoverable
        }


def sanitize_text(text: str) -> str:
    """
    Remove control characters that are not valid in XML/DOCX.
    Keeps normal whitespace (space, tab, newline, carriage return).
    """
    if not text:
        return ""
    
    result = []
    for char in text:
        code = ord(char)
        # Valid XML chars: #x9 | #xA | #xD | [#x20-#xD7FF] | [#xE000-#xFFFD]
        if code == 0x9 or code == 0xA or code == 0xD or (code >= 0x20 and code <= 0xD7FF) or (code >= 0xE000 and code <= 0xFFFD):
            result.append(char)
        elif code < 0x20:
            result.append(' ')
    
    return ''.join(result)


def get_rgb_from_color(color_value) -> tuple:
    """Convert PyMuPDF color value to RGB tuple."""
    if color_value is None:
        return (0, 0, 0)
    
    if isinstance(color_value, (list, tuple)):
        if len(color_value) == 3:
            return tuple(int(c * 255) for c in color_value)
        elif len(color_value) == 1:
            gray = int(color_value[0] * 255)
            return (gray, gray, gray)
        elif len(color_value) == 4:
            c, m, y, k = color_value
            r = int(255 * (1 - c) * (1 - k))
            g = int(255 * (1 - m) * (1 - k))
            b = int(255 * (1 - y) * (1 - k))
            return (r, g, b)
    elif isinstance(color_value, (int, float)):
        if isinstance(color_value, float):
            gray = int(color_value * 255)
            return (gray, gray, gray)
        else:
            if color_value == 0:
                return (0, 0, 0)
            r = (color_value >> 16) & 0xFF
            g = (color_value >> 8) & 0xFF
            b = color_value & 0xFF
            return (r, g, b)
    
    return (0, 0, 0)


def extract_images_from_page(page, page_num: int, temp_dir: str) -> list:
    """Extract all images from a PDF page and save them to temp files."""
    images = []
    
    try:
        image_list = page.get_images(full=True)
        
        for img_index, img_info in enumerate(image_list):
            try:
                xref = img_info[0]
                base_image = page.parent.extract_image(xref)
                if not base_image:
                    continue
                
                image_bytes = base_image.get("image")
                image_ext = base_image.get("ext", "png")
                
                if not image_bytes:
                    continue
                
                # Convert CMYK to RGB if needed
                if HAS_PIL and image_ext in ["jpeg", "jpg"]:
                    try:
                        img = Image.open(io.BytesIO(image_bytes))
                        if img.mode == "CMYK":
                            img = img.convert("RGB")
                            buffer = io.BytesIO()
                            img.save(buffer, format="PNG")
                            image_bytes = buffer.getvalue()
                            image_ext = "png"
                    except Exception as pil_err:
                        print(json.dumps({"warning": f"PIL conversion failed: {pil_err}"}), file=sys.stderr)
                
                image_filename = f"page{page_num}_img{img_index}.{image_ext}"
                image_path = os.path.join(temp_dir, image_filename)
                
                with open(image_path, "wb") as img_file:
                    img_file.write(image_bytes)
                
                img_rects = page.get_image_rects(xref)
                if img_rects:
                    bbox = img_rects[0]
                    images.append((image_path, bbox, base_image.get("width", 100), base_image.get("height", 100)))
                else:
                    images.append((image_path, None, base_image.get("width", 100), base_image.get("height", 100)))
                    
            except Exception as img_err:
                print(json.dumps({"warning": f"Failed to extract image {img_index}: {img_err}"}), file=sys.stderr)
                continue
    
    except Exception as e:
        print(json.dumps({"warning": f"Image extraction error on page {page_num}: {e}"}), file=sys.stderr)
    
    images.sort(key=lambda x: x[1][1] if x[1] else float('inf'))
    return images


def add_image_to_doc(doc, image_path: str, width: int = None, height: int = None, max_width_inches: float = 6.0):
    """Add an image to the document with appropriate sizing."""
    try:
        if width and height:
            aspect_ratio = height / width
            img_width = min(max_width_inches, width / 96)
            img_height = img_width * aspect_ratio
        else:
            img_width = max_width_inches / 2
            img_height = None
        
        para = doc.add_paragraph()
        run = para.add_run()
        
        if img_height:
            run.add_picture(image_path, width=Inches(img_width), height=Inches(img_height))
        else:
            run.add_picture(image_path, width=Inches(img_width))
        
        return True
    except Exception as e:
        print(json.dumps({"warning": f"Failed to add image: {e}"}), file=sys.stderr)
        return False


def add_formatted_paragraph_with_breaks(doc, runs: list, page_avg_size: float):
    """Add a paragraph to the document with formatted runs and line breaks."""
    if not runs:
        return
    
    full_text = ''.join(r[0] for r in runs if not r[6]).strip()
    if not full_text:
        return
    
    para = doc.add_paragraph()
    
    first_non_break = next((r for r in runs if not r[6]), None)
    first_size = first_non_break[2] if first_non_break else 12
    is_heading = first_size > page_avg_size * 1.3
    
    for run_data in runs:
        text, font_name, font_size, color, is_bold, is_italic, is_line_break = run_data
        
        if is_line_break:
            run = para.add_run()
            run.add_break(WD_BREAK.LINE)
            continue
        
        if not text:
            continue
        
        run = para.add_run(text)
        run.font.size = Pt(font_size)
        
        try:
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        except:
            run.font.name = 'Arial'
        
        if color != (0, 0, 0):
            try:
                run.font.color.rgb = RGBColor(color[0], color[1], color[2])
            except:
                pass
        
        if is_bold or is_heading:
            run.font.bold = True
        if is_italic:
            run.font.italic = True


def extract_text_to_docx(pdf_path: str, output_path: str) -> dict:
    """
    Extract text from PDF and create a DOCX with preserved formatting.
    Fallback method when other converters fail.
    """
    if not HAS_PYTHON_DOCX:
        return {
            "success": False,
            "error": "python-docx not installed. Install with: pip install python-docx",
            "method": "rich_text_extraction"
        }
    
    if not HAS_PYMUPDF:
        return {
            "success": False,
            "error": "PyMuPDF not installed. Install with: pip install PyMuPDF",
            "method": "rich_text_extraction"
        }
    
    temp_dir = tempfile.mkdtemp(prefix="pdf_images_")
    
    try:
        print(json.dumps({"info": "Using rich text extraction fallback..."}), file=sys.stderr)
        
        pdf_doc = fitz.open(pdf_path)
        total_pages = len(pdf_doc)
        
        doc = Document()
        
        print(json.dumps({"info": f"Extracting from {total_pages} pages..."}), file=sys.stderr)
        
        total_images = 0
        
        for page_num in range(total_pages):
            page = pdf_doc[page_num]
            page_height = page.rect.height
            
            page_images = extract_images_from_page(page, page_num, temp_dir)
            total_images += len(page_images)
            
            text_dict = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)
            blocks = text_dict.get("blocks", [])
            
            # Calculate average font size
            all_sizes = []
            for block in blocks:
                if block.get("type") == 0:
                    for line in block.get("lines", []):
                        for span in line.get("spans", []):
                            all_sizes.append(span.get("size", 12))
            
            page_avg_size = sum(all_sizes) / len(all_sizes) if all_sizes else 12
            
            # Combine content items sorted by position
            content_items = []
            
            for block in blocks:
                if block.get("type") == 0:
                    bbox = block.get("bbox", [0, 0, 0, 0])
                    content_items.append(("text", bbox[1], block))
            
            for img_path, img_bbox, img_width, img_height in page_images:
                y_pos = img_bbox[1] if img_bbox else page_height
                content_items.append(("image", y_pos, (img_path, img_width, img_height)))
            
            content_items.sort(key=lambda x: x[1])
            
            # Process content
            for item_type, y_pos, item_data in content_items:
                if item_type == "text":
                    block = item_data
                    lines = block.get("lines", [])
                    if not lines:
                        continue
                    
                    current_para_runs = []
                    last_y1 = None
                    last_size = None
                    
                    for line in lines:
                        spans = line.get("spans", [])
                        if not spans:
                            continue
                        
                        line_bbox = line.get("bbox", [0, 0, 0, 0])
                        line_y0 = line_bbox[1]
                        line_y1 = line_bbox[3]
                        line_height = line_y1 - line_y0 if line_y1 > line_y0 else 12
                        
                        start_new_para = False
                        add_line_break = False
                        
                        if last_y1 is not None:
                            gap = line_y0 - last_y1
                            if gap > line_height * 1.0:
                                start_new_para = True
                            elif gap > line_height * 0.2:
                                add_line_break = True
                        
                        first_span = spans[0]
                        current_size = first_span.get("size", 12)
                        if last_size is not None and abs(current_size - last_size) > 3:
                            start_new_para = True
                        
                        if start_new_para and current_para_runs:
                            add_formatted_paragraph_with_breaks(doc, current_para_runs, page_avg_size)
                            current_para_runs = []
                            add_line_break = False
                        
                        if add_line_break and current_para_runs:
                            last_run = current_para_runs[-1]
                            current_para_runs.append(('\n', last_run[1], last_run[2], last_run[3], last_run[4], last_run[5], True))
                        
                        for span in spans:
                            text = span.get("text", "")
                            if not text:
                                continue
                            
                            text = sanitize_text(text)
                            if not text:
                                continue
                            
                            font_name = span.get("font", "Arial")
                            font_size = span.get("size", 12)
                            color = get_rgb_from_color(span.get("color"))
                            flags = span.get("flags", 0)
                            is_bold = bool(flags & (1 << 4))
                            is_italic = bool(flags & (1 << 1))
                            
                            if '+' in font_name:
                                font_name = font_name.split('+', 1)[1]
                            
                            current_para_runs.append((text, font_name, font_size, color, is_bold, is_italic, False))
                        
                        last_y1 = line_y1
                        last_size = current_size
                    
                    if current_para_runs:
                        add_formatted_paragraph_with_breaks(doc, current_para_runs, page_avg_size)
                
                elif item_type == "image":
                    img_path, img_width, img_height = item_data
                    add_image_to_doc(doc, img_path, img_width, img_height)
            
            # Page break between pages
            if page_num < total_pages - 1:
                doc.add_page_break()
            
            if (page_num + 1) % 10 == 0 or (page_num + 1) == total_pages:
                progress = int((page_num + 1) / total_pages * 100)
                print(json.dumps({"info": f"Progress: {page_num + 1}/{total_pages} ({progress}%)"}), file=sys.stderr)
        
        pdf_doc.close()
        doc.save(output_path)
        
        file_size = os.path.getsize(output_path)
        print(json.dumps({"info": f"Rich text extraction complete: {total_pages} pages, {total_images} images, {file_size} bytes"}), file=sys.stderr)
        
        return {
            "success": True,
            "outputPath": output_path,
            "method": "rich_text_extraction"
        }
        
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(json.dumps({"error": f"Text extraction error: {error_details}"}), file=sys.stderr)
        return {
            "success": False,
            "error": f"Text extraction failed: {str(e)}",
            "method": "rich_text_extraction"
        }
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)


def convert_pdf_to_docx(pdf_path: str, output_path: str) -> dict:
    """
    Convert PDF to DOCX using the best available method.
    
    Tries methods in order of quality:
    1. LibreOffice headless (best layout preservation)
    2. pdf2docx (good for most PDFs)
    3. Rich text extraction (fallback)
    """
    print(json.dumps({"info": "Starting PDF to DOCX conversion (hybrid approach)..."}), file=sys.stderr)
    
    # Verify input file exists
    if not os.path.exists(pdf_path):
        return {
            "success": False,
            "error": f"Input PDF file not found: {pdf_path}"
        }
    
    methods_tried = []
    
    # Method 1: Try LibreOffice (best quality)
    print(json.dumps({"info": "Attempting Method 1: LibreOffice..."}), file=sys.stderr)
    result = convert_with_libreoffice(pdf_path, output_path)
    methods_tried.append(f"LibreOffice: {result.get('error', 'success') if not result['success'] else 'success'}")
    
    if result["success"]:
        print(json.dumps({"info": "✓ LibreOffice conversion successful"}), file=sys.stderr)
        return result
    else:
        print(json.dumps({"warning": f"LibreOffice failed: {result.get('error', 'unknown')}"}), file=sys.stderr)
    
    # Method 2: Try pdf2docx
    print(json.dumps({"info": "Attempting Method 2: pdf2docx..."}), file=sys.stderr)
    result = convert_with_pdf2docx(pdf_path, output_path)
    methods_tried.append(f"pdf2docx: {result.get('error', 'success') if not result['success'] else 'success'}")
    
    if result["success"]:
        print(json.dumps({"info": "✓ pdf2docx conversion successful"}), file=sys.stderr)
        return result
    else:
        print(json.dumps({"warning": f"pdf2docx failed: {result.get('error', 'unknown')}"}), file=sys.stderr)
    
    # Method 3: Rich text extraction fallback
    print(json.dumps({"info": "Attempting Method 3: Rich text extraction..."}), file=sys.stderr)
    result = extract_text_to_docx(pdf_path, output_path)
    methods_tried.append(f"Rich text: {result.get('error', 'success') if not result['success'] else 'success'}")
    
    if result["success"]:
        print(json.dumps({"info": "✓ Rich text extraction successful"}), file=sys.stderr)
        return result
    
    # All methods failed
    return {
        "success": False,
        "error": f"All conversion methods failed. Tried: {'; '.join(methods_tried)}"
    }


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print(json.dumps({
            "success": False,
            "error": "Usage: pdf_to_docx.py <pdf_file_path> <output_docx_path>"
        }))
        sys.exit(1)
    
    pdf_path = sys.argv[1]
    output_path = sys.argv[2]
    
    result = convert_pdf_to_docx(pdf_path, output_path)
    print(json.dumps(result))
